"""Document loading utilities for PDF, DOCX, and TXT resumes."""

from __future__ import annotations

import logging
from pathlib import Path

from parser.errors import DocumentLoadError, UnsupportedFileTypeError

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}


def load_document(path: str | Path) -> str:
    """Extract text from a PDF, DOCX, or TXT file."""

    file_path = Path(path)
    suffix = file_path.suffix.lower()
    logger.info("Loading document: %s", file_path)

    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{suffix}'. Supported types: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    if not file_path.exists():
        raise DocumentLoadError(f"File does not exist: {file_path}")

    try:
        if suffix == ".txt":
            return file_path.read_text(encoding="utf-8")
        if suffix == ".docx":
            return _load_docx(file_path)
        if suffix == ".pdf":
            return _load_pdf(file_path)
    except DocumentLoadError:
        raise
    except Exception as exc:  # pragma: no cover - defensive wrapper
        raise DocumentLoadError(f"Failed to load {file_path}: {exc}") from exc

    raise UnsupportedFileTypeError(f"Unsupported file type: {suffix}")


def _load_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise DocumentLoadError("DOCX support requires python-docx to be installed.") from exc

    document = Document(str(path))
    text_blocks = [_normalize_text(paragraph.text) for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text_blocks.extend(_normalize_text(paragraph.text) for paragraph in cell.paragraphs)
    text = "\n".join(block for block in text_blocks if block)
    if not text.strip():
        raise DocumentLoadError(f"No text found in DOCX: {path}")
    return text


def _normalize_text(text: str) -> str:
    return " ".join(text.split())


def _load_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise DocumentLoadError("PDF support requires pypdf to be installed.") from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(page.strip() for page in pages if page.strip())
    if not text.strip():
        raise DocumentLoadError(f"No text found in PDF: {path}")
    return text
